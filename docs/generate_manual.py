# -*- coding: utf-8 -*-
"""
Génère le Manuel d'utilisation TEL & CASH au format .docx (python-docx).
Document de livraison client : parcours acheteur, back-office, opérations, annexe technique.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from PIL import Image, ImageOps

# Dossier des captures et cache des images encadrées
CAP_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "captures")
FRAME_DIR = os.path.join(CAP_DIR, ".framed")
os.makedirs(FRAME_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Palette de marque
# ---------------------------------------------------------------------------
NAVY   = "16233F"   # bleu nuit (page de garde, titres)
BRAND  = "2F6BFF"   # bleu primaire TEL & CASH
BRAND2 = "1E4FD6"   # bleu foncé
INK    = "1F2937"   # texte principal
GRAY   = "6B7280"   # texte secondaire
LGRAY  = "9AA3B2"   # gris clair
LINE   = "E3E8F0"   # filets

# Encadrés
INFO_BG,   INFO_BD   = "EAF1FF", "2F6BFF"
WARN_BG,   WARN_BD   = "FFF6E5", "E8A13A"
DANGER_BG, DANGER_BD = "FDECEC", "D64545"
OK_BG,     OK_BD     = "E9F7EF", "2FA36B"
TIP_BG,    TIP_BD    = "F1ECFB", "7C5CD6"


def C(hexstr):
    return RGBColor.from_string(hexstr)


# ---------------------------------------------------------------------------
# Bas-niveau XML
# ---------------------------------------------------------------------------
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def cell_vertical_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcPr.append(va)


def set_left_border(cell, hex_color, size=24):
    """Grosse bordure gauche colorée (pour les encadrés)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "start", "end"):
        e = OxmlElement(f"w:{edge}")
        if edge == "start":
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(size))
            e.set(qn("w:color"), hex_color)
        else:
            e.set(qn("w:val"), "nil")
        borders.append(e)
    tcPr.append(borders)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tblPr.append(borders)


def keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    e = OxmlElement("w:keepNext")
    pPr.append(e)


def add_bottom_rule(paragraph, hex_color=BRAND, size=8):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), hex_color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_field(paragraph, code):
    run = paragraph.add_run()
    r = run._r
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = code
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r.append(f1); r.append(it); r.append(f2)
    return run


# ---------------------------------------------------------------------------
# Styles du document
# ---------------------------------------------------------------------------
def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = C(INK)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(19)
    h1.font.bold = True
    h1.font.color.rgb = C(NAVY)
    h1.paragraph_format.space_before = Pt(4)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = C(BRAND2)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = C(INK)
    h3.paragraph_format.space_before = Pt(9)
    h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.keep_with_next = True


# ---------------------------------------------------------------------------
# Helpers de contenu
# ---------------------------------------------------------------------------
def _emit_runs(paragraph, text):
    """Transforme **gras** en runs gras."""
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if seg == "":
            continue
        run = paragraph.add_run(seg)
        if i % 2 == 1:
            run.bold = True


def para(doc, text="", size=None, color=None, italic=False, align=None,
         space_after=6, space_before=0, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    if text:
        _emit_runs(p, text)
        for run in p.runs:
            if size:  run.font.size = Pt(size)
            if color: run.font.color.rgb = C(color)
            if italic: run.italic = True
            if bold:  run.bold = True
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if level:
        p.paragraph_format.left_indent = Inches(0.5 + 0.25 * level)
    _emit_runs(p, text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    _emit_runs(p, text)
    return p


def h1(doc, text, page_break=True):
    if page_break:
        doc.add_page_break()
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.add_run("").font.size = Pt(pts)
    return p


def _framed(filename):
    """Ajoute un fin cadre gris + marge blanche autour de la capture."""
    src = os.path.join(CAP_DIR, filename)
    dst = os.path.join(FRAME_DIR, filename)
    im = Image.open(src).convert("RGB")
    im = ImageOps.expand(im, border=1, fill=(227, 232, 240))   # filet LINE
    im = ImageOps.expand(im, border=6, fill=(255, 255, 255))   # marge blanche
    im = ImageOps.expand(im, border=1, fill=(214, 221, 232))   # contour externe
    im.save(dst)
    return dst


def figure(doc, filename, caption, width_in=6.15):
    path = _framed(filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(path, width=Inches(width_in))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = C(GRAY)
    return p


def callout(doc, title, body_lines, kind="info"):
    themes = {
        "info":   (INFO_BG, INFO_BD, "ℹ", "À SAVOIR"),
        "warning":(WARN_BG, WARN_BD, "▲", "ATTENTION"),
        "danger": (DANGER_BG, DANGER_BD, "✕", "IMPORTANT"),
        "ok":     (OK_BG, OK_BD, "✓", "BONNE PRATIQUE"),
        "tip":    (TIP_BG, TIP_BD, "★", "ASTUCE"),
    }
    bg, bd, icon, default_kicker = themes[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_left_border(cell, bd, size=28)
    set_cell_margins(cell, top=90, bottom=90, left=170, right=140)

    # Titre de l'encadré
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(3)
    r = p0.add_run(f"{icon}  {title if title else default_kicker}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = C(bd)

    if isinstance(body_lines, str):
        body_lines = [body_lines]
    for i, line in enumerate(body_lines):
        bp = cell.add_paragraph()
        bp.paragraph_format.space_after = Pt(1)
        _emit_runs(bp, line)
        for run in bp.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = C(INK)
    spacer(doc, 4)
    return table


def data_table(doc, headers, rows, widths=None, header_fill=NAVY,
               zebra="F5F7FB", font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # bordures fines gris clair
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), LINE)
        borders.append(e)
    tblPr.append(borders)

    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_bg(hdr[i], header_fill)
        set_cell_margins(hdr[i])
        cell_vertical_center(hdr[i])
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(htext)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = C("FFFFFF")

    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i])
            cell_vertical_center(cells[i])
            if ridx % 2 == 1:
                set_cell_bg(cells[i], zebra)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            _emit_runs(p, str(val))
            for run in p.runs:
                run.font.size = Pt(font_size)

    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    spacer(doc, 5)
    return table


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    r = run._r
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = 'TOC \\o "1-3" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t  = OxmlElement("w:t"); t.text = "Le sommaire se génère à l'ouverture (sinon : clic droit ▸ Mettre à jour les champs)."
    rr = OxmlElement("w:r"); rr.append(t)
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    r.append(f1); r.append(it); r.append(f2); r.append(rr); r.append(f3)


def enable_update_fields(doc):
    settings = doc.settings.element
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    settings.append(uf)


# ---------------------------------------------------------------------------
# Page de garde
# ---------------------------------------------------------------------------
def build_cover(doc):
    # Bandeau supérieur bleu nuit
    band = doc.add_table(rows=1, cols=1)
    no_table_borders(band)
    c = band.rows[0].cells[0]
    set_cell_bg(c, NAVY)
    set_cell_margins(c, top=460, bottom=460, left=360, right=360)

    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("TEL & CASH")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = C("FFFFFF")

    p2 = c.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    r2 = p2.add_run("Smartphones reconditionnés  ·  Boutique en ligne")
    r2.font.size = Pt(11)
    r2.font.color.rgb = C("AFC4FF")

    # Titre principal
    spacer(doc, 26)
    t = doc.add_paragraph()
    t.paragraph_format.space_after = Pt(2)
    rt = t.add_run("Manuel d’utilisation")
    rt.bold = True
    rt.font.size = Pt(34)
    rt.font.color.rgb = C(NAVY)

    st = doc.add_paragraph()
    add_bottom_rule(st, BRAND, size=14)
    rst = st.add_run("Guide complet de la plateforme : boutique client & back-office d’administration")
    rst.font.size = Pt(13)
    rst.font.color.rgb = C(GRAY)

    spacer(doc, 30)

    # Bloc méta
    meta = doc.add_table(rows=0, cols=2)
    no_table_borders(meta)
    rows_meta = [
        ("Livrable", "Documentation fonctionnelle & technique"),
        ("Destinataire", "Exploitant du site TEL & CASH"),
        ("Version du document", "1.0"),
        ("Date", "7 juillet 2026"),
        ("Périmètre", "Parcours d’achat · Gestion des ventes, retours & encaissements · Suivi des dossiers · Annexe technique"),
        ("Confidentialité", "Document interne — diffusion restreinte"),
    ]
    for k, v in rows_meta:
        row = meta.add_row().cells
        pk = row[0].paragraphs[0]
        pk.paragraph_format.space_after = Pt(4)
        rk = pk.add_run(k.upper())
        rk.bold = True
        rk.font.size = Pt(8.5)
        rk.font.color.rgb = C(BRAND)
        pv = row[1].paragraphs[0]
        pv.paragraph_format.space_after = Pt(4)
        rv = pv.add_run(v)
        rv.font.size = Pt(10.5)
        rv.font.color.rgb = C(INK)
        row[0].width = Inches(2.1)
        row[1].width = Inches(4.4)


# ---------------------------------------------------------------------------
# En-tête / pied de page
# ---------------------------------------------------------------------------
def build_header_footer(section):
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("TEL & CASH · Manuel d’utilisation")
    hr.font.size = Pt(8)
    hr.font.color.rgb = C(LGRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    fr = fp.add_run("© 2026 TEL & CASH — Confidentiel")
    fr.font.size = Pt(8)
    fr.font.color.rgb = C(LGRAY)
    fp.add_run("\t")
    pr = fp.add_run("Page ")
    pr.font.size = Pt(8)
    pr.font.color.rgb = C(GRAY)
    fld = add_field(fp, "PAGE")
    fld.font.size = Pt(8)
    fld.font.color.rgb = C(GRAY)


# ===========================================================================
# CONSTRUCTION DU DOCUMENT
# ===========================================================================
doc = Document()
setup_styles(doc)

sec = doc.sections[0]
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(1.8)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)
build_header_footer(sec)

# ---- Page de garde ----
build_cover(doc)

# ---- Sommaire ----
doc.add_page_break()
tt = doc.add_paragraph()
add_bottom_rule(tt, BRAND, size=10)
rtt = tt.add_run("Sommaire")
rtt.bold = True
rtt.font.size = Pt(20)
rtt.font.color.rgb = C(NAVY)
spacer(doc, 4)
add_toc(doc)


# ===========================================================================
# PARTIE I — PRÉSENTATION GÉNÉRALE
# ===========================================================================
h1(doc, "Partie I — Présentation générale")

h2(doc, "1. Objet du document")
para(doc, "Ce manuel décrit le fonctionnement complet de la plateforme e-commerce **TEL & CASH**, "
          "spécialisée dans la vente de smartphones reconditionnés. Il s’adresse à deux publics :")
bullet(doc, "**L’équipe d’exploitation** (gérant, vendeurs, service après-vente) qui utilise l’espace "
            "d’administration pour gérer les ventes, les retours, les encaissements et le suivi des dossiers clients.")
bullet(doc, "**Le responsable technique** chargé de la maintenance, qui trouvera en annexe l’architecture, "
            "les intégrations (paiement, fournisseur) et la configuration.")
para(doc, "Il peut être remis tel quel comme documentation de référence, servir de support de formation "
          "pour un nouveau collaborateur, ou de base à une reprise technique du projet.")

callout(doc, "Comment lire ce manuel",
        ["La **Partie II** suit le parcours d’un client, de la visite jusqu’au suivi de sa commande.",
         "La **Partie III** détaille chaque écran de l’espace d’administration.",
         "La **Partie IV** regroupe des routines et check-lists opérationnelles prêtes à l’emploi.",
         "La **Partie V** est l’annexe technique (hébergement, base de données, intégrations, variables)."],
        kind="info")

h3(doc, "Conventions & pictogrammes")
para(doc, "Les encadrés colorés attirent l’attention sur un point précis :")
bullet(doc, "**ℹ À savoir** — information de contexte utile.")
bullet(doc, "**▲ Attention** — point de vigilance à ne pas négliger.")
bullet(doc, "**✕ Important** — action sensible ou irréversible.")
bullet(doc, "**✓ Bonne pratique** — recommandation pour un fonctionnement optimal.")

h2(doc, "2. Vue d’ensemble de la plateforme")
para(doc, "TEL & CASH est un site marchand complet. Il se compose de deux univers complémentaires :")
data_table(
    doc,
    ["Univers", "Adresse", "Public", "Rôle"],
    [
        ["La **boutique**", "Pages publiques du site", "Visiteurs & clients", "Découvrir le catalogue, commander, payer, suivre ses commandes, demander un retour."],
        ["L’**administration**", "Section « /admin »", "Équipe TEL & CASH", "Piloter les ventes, expéditions, retours, remboursements, prix, stock et clients."],
    ],
    widths=[1.4, 1.5, 1.5, 2.4],
)
para(doc, "Le site fonctionne en **vente à la commande** (« sell-to-order ») : les produits affichés sont "
          "approvisionnés auprès d’un fournisseur au moment de la vente. Le client ne voit donc jamais de "
          "notion de « stock » ou de « rupture » ; il voit un catalogue de modèles disponibles à la commande.")

h2(doc, "3. Concepts clés")
para(doc, "Quelques notions structurantes reviennent tout au long du manuel :")

h3(doc, "Le reconditionné et les grades")
para(doc, "Chaque téléphone est un produit d’occasion remis en état, classé par **grade** correspondant à "
          "son état esthétique. Trois grades sont proposés aux clients :")
data_table(
    doc,
    ["Grade", "État", "Description client"],
    [
        ["**A**", "Excellent", "Aspect quasi neuf, micro-rayures éventuelles à peine visibles."],
        ["**B**", "Bon", "Traces d’usage légères, rayures visibles de près."],
        ["**C**", "Correct", "Marques d’usage plus marquées, pleinement fonctionnel."],
    ],
    widths=[1.0, 1.5, 4.0],
)
callout(doc, "Grades D et E exclus",
        ["Les grades inférieurs (**D**, **E**) ne sont **jamais** proposés à la vente sur la boutique : "
         "ils sont automatiquement écartés du catalogue. Seuls A, B et C sont visibles côté client."],
        kind="info")

h3(doc, "La notion de variante")
para(doc, "Un même modèle (ex. « iPhone 13 ») se décline en **variantes**. Une variante est une combinaison "
          "unique de trois critères :")
bullet(doc, "**Stockage** : 64 Go, 128 Go, 256 Go, 512 Go…")
bullet(doc, "**Grade** : A, B ou C.")
bullet(doc, "**Couleur** : noir, blanc, bleu, etc.")
para(doc, "Chaque variante possède son **propre prix** et ses **propres photos**. Sur la fiche produit, "
          "le client compose sa variante en sélectionnant successivement stockage, grade puis couleur.")

h3(doc, "Cycle de vie d’une vente")
para(doc, "Une vente passe par une suite d’états maîtrisés, de la commande du client jusqu’à la livraison. "
          "Ce cycle est détaillé en Partie III ; en résumé :")
para(doc, "**Payée → Commandée chez le fournisseur → Expédiée → Livrée**, avec les branches "
          "**Annulée / Remboursée** et **Litige** en cas d’incident.", italic=True)

h2(doc, "4. Rôles et accès")
data_table(
    doc,
    ["Profil", "Accès", "Peut faire"],
    [
        ["Visiteur", "Boutique publique", "Parcourir le catalogue, remplir un panier. Doit créer un compte pour payer."],
        ["Client", "Boutique + espace compte", "Commander, payer, suivre ses commandes, demander un retour, gérer sa fidélité."],
        ["Administrateur", "Boutique + espace « /admin »", "Toutes les opérations de gestion (ventes, retours, encaissements, catalogue, clients)."],
    ],
    widths=[1.3, 2.1, 3.4],
)
callout(doc, "Accès administrateur",
        ["L’accès à l’espace d’administration est **strictement réservé** aux comptes disposant du rôle "
         "« administrateur ». Toute tentative d’accès sans ce rôle est bloquée automatiquement. "
         "Voir l’annexe technique pour attribuer le rôle admin à un compte."],
        kind="warning")


# ===========================================================================
# PARTIE II — PARCOURS CLIENT
# ===========================================================================
h1(doc, "Partie II — Le parcours client (côté acheteur)")
para(doc, "Cette partie suit un client de bout en bout : de sa première visite jusqu’au suivi de sa commande "
          "et, le cas échéant, à une demande de retour. Elle permet de comprendre ce que vit l’acheteur et de "
          "répondre à ses questions au téléphone ou par e-mail.")

h2(doc, "5. La page d’accueil")
para(doc, "La page d’accueil présente l’enseigne et oriente le visiteur vers le catalogue. Elle enchaîne "
          "plusieurs sections :")
bullet(doc, "Une **bannière principale** (hero) avec un appel à l’action vers la boutique.")
bullet(doc, "Les **marques populaires** (Apple, Samsung, Xiaomi…) menant directement au catalogue filtré.")
bullet(doc, "Les **meilleures offres** et **best-sellers**.")
bullet(doc, "Des sections pédagogiques : pourquoi le reconditionné, explication des grades, garantie, "
            "fonctionnement en quelques étapes, avis, FAQ, inscription à la newsletter.")
para(doc, "Chaque marque ou produit mis en avant renvoie vers la page correspondante du catalogue.")
figure(doc, "01-accueil.png",
       "Figure 1 — La page d’accueil : bannière principale, mise en avant des marques et réassurance "
       "(la bannière de consentement cookies apparaît en bas à gauche).")

h2(doc, "6. Naviguer dans le catalogue")
para(doc, "Le catalogue (« Nos téléphones ») liste les modèles disponibles. Le client affine sa recherche "
          "grâce à des filtres qui se combinent librement :")
data_table(
    doc,
    ["Filtre", "Options"],
    [
        ["Marque", "Apple, Samsung, Xiaomi, Google… (selon le catalogue)"],
        ["Grade / état", "A (Excellent), B (Bon), C (Correct)"],
        ["Stockage", "64, 128, 256, 512 Go…"],
        ["Prix", "Curseur de fourchette min–max"],
        ["Recherche", "Champ texte tolérant aux accents et espaces (« galaxy s22 » = « Galaxy S22 »)"],
    ],
    widths=[1.6, 4.9],
)
para(doc, "Les résultats se **trient** par pertinence (par défaut), prix croissant, prix décroissant ou nom. "
          "Les filtres sont mémorisés dans l’adresse de la page : un client peut donc partager un lien "
          "de recherche déjà filtré.")
para(doc, "Chaque téléphone est présenté comme une **carte de modèle** unique regroupant toutes ses "
          "variantes, avec une photo, la marque, le modèle et un prix **« à partir de »** (le prix de la "
          "variante la moins chère disponible).")
figure(doc, "02-catalogue.png",
       "Figure 2 — Le catalogue et ses filtres (marque, grade, stockage, prix), avec le tri et la recherche.")
callout(doc, "Pas d’affichage de stock",
        ["Conformément au fonctionnement en vente à la commande, **aucune mention de stock ou de rupture** "
         "n’apparaît côté client. Tout produit visible est commandable. La disponibilité réelle est vérifiée "
         "en coulisses au moment du paiement (voir §11)."],
        kind="info")

h3(doc, "Les accessoires")
para(doc, "Le catalogue comporte aussi une section **accessoires** (câbles, chargeurs, coques, protections…). "
          "Ce sont des produits simples, sans grade ni stockage : le client voit une photo, un nom et un prix "
          "unique, et peut les ajouter directement au panier. Les filtres se limitent alors à la marque, "
          "au type d’accessoire et au prix.")

h2(doc, "7. La fiche produit")
para(doc, "En cliquant sur un modèle, le client arrive sur la **fiche produit**. Il y compose sa variante :")
numbered(doc, "**Stockage** — il choisit la capacité souhaitée.")
numbered(doc, "**Grade / état** — il choisit A, B ou C ; le prix se met à jour.")
numbered(doc, "**Couleur** — il choisit la couleur ; la photo principale s’adapte en temps réel.")
para(doc, "Les combinaisons non disponibles sont grisées et non sélectionnables. Le **prix affiché** "
          "correspond exactement à la variante choisie (avec, le cas échéant, un prix barré indiquant une "
          "promotion).")
para(doc, "La fiche affiche également les **garanties et informations de réassurance** : garantie légale "
          "24 mois, produit testé et certifié en France, délai de livraison indicatif (5 à 10 jours ouvrés), "
          "possibilité de retour, moyens de paiement (Visa, Mastercard, paiement en plusieurs fois), ainsi "
          "que les caractéristiques techniques du modèle.")
para(doc, "Le bouton **« Ajouter au panier »** est actif dès qu’une variante valide est sélectionnée. "
          "Un message de confirmation s’affiche à l’ajout.")
figure(doc, "03-fiche-produit.png",
       "Figure 3 — La fiche produit : composition de la variante, prix, garanties et ajout au panier.")

h2(doc, "8. Le panier")
para(doc, "Le panier récapitule les articles sélectionnés. Le client peut :")
bullet(doc, "Modifier les **quantités** (limite de 10 unités par ligne).")
bullet(doc, "**Supprimer** un article.")
bullet(doc, "Revenir à la boutique pour continuer ses achats.")
para(doc, "Le panier affiche le **sous-total**, les **frais de livraison** (9,90 € par défaut) et le "
          "**total TTC**. Un rappel de la possibilité de paiement en plusieurs fois est présenté.")
callout(doc, "Panier conservé automatiquement",
        ["Le panier est **mémorisé dans le navigateur** du visiteur : il survit à un rafraîchissement de la "
         "page ou à une fermeture de l’onglet. Lorsqu’un visiteur se connecte, son panier local est "
         "**fusionné** avec son panier enregistré sur le compte."],
        kind="info")
figure(doc, "04-panier.png",
       "Figure 4 — Le panier : articles et quantités, frais de livraison et total TTC.")

h2(doc, "9. Créer un compte et se connecter")
para(doc, "La **création de compte** demande un nom complet, un e-mail et un mot de passe (6 caractères "
          "minimum). La connexion est immédiate — il n’y a pas d’e-mail de confirmation à valider. "
          "La connexion se fait ensuite par e-mail + mot de passe, avec un lien « mot de passe oublié ».")
callout(doc, "Un compte est obligatoire pour payer",
        ["Le passage de commande **sans compte** (achat « invité ») est actuellement **désactivé**. "
         "Un visiteur qui clique sur « Passer au paiement » sans être connecté est redirigé vers la page de "
         "connexion, puis ramené automatiquement à son paiement une fois identifié."],
        kind="warning")
figure(doc, "05-connexion.png",
       "Figure 5 — La page de connexion (avec accès à la création de compte et au mot de passe oublié).")

h2(doc, "10. Passer commande (checkout)")
para(doc, "Le paiement se déroule en trois étapes claires :")
h3(doc, "Étape 1 — Récapitulatif du panier")
para(doc, "Le client vérifie les articles et leur montant.")
h3(doc, "Étape 2 — Livraison")
para(doc, "Il renseigne son **adresse de livraison** : nom, prénom, pays (France, Belgique, Luxembourg, "
          "Suisse), adresse, code postal, ville et téléphone. Le mode de livraison proposé est la "
          "**livraison à domicile avec suivi** (5 à 10 jours ouvrés, 9,90 €).")
h3(doc, "Étape 3 — Paiement")
para(doc, "Le client relit son adresse, **accepte les conditions générales de vente** (case obligatoire, avec "
          "liens vers les CGV et la politique de retour), puis clique sur **« Procéder au paiement »**.")
para(doc, "Il est alors redirigé vers la **page de paiement sécurisée Stripe** (en français), où il saisit "
          "ses coordonnées bancaires. Les moyens disponibles incluent carte bancaire (Visa, Mastercard), "
          "portefeuilles mobiles et paiement en plusieurs fois (selon la configuration Stripe).")

callout(doc, "Que se passe-t-il après le paiement ?",
        ["Une fois le paiement validé, Stripe notifie automatiquement le site. Celui-ci **confirme la "
         "commande**, **vérifie la disponibilité** auprès du fournisseur, vide le panier, attribue les points "
         "de fidélité et envoie l’**e-mail de confirmation** au client (ainsi qu’une alerte à l’équipe). "
         "Une **facture PDF** est générée automatiquement."],
        kind="info")

callout(doc, "Sécurité anti-survente",
        ["Si, au moment exact du paiement, le produit s’avère indisponible chez le fournisseur, la commande "
         "est **remboursée automatiquement** et le client en est informé par e-mail. Aucune commande ne peut "
         "donc être encaissée sans possibilité d’approvisionnement."],
        kind="ok")

h2(doc, "11. Confirmation et suivi de commande")
para(doc, "Après paiement, le client atterrit sur une page **« Commande confirmée ! »** avec sa référence. "
          "Deux boutons lui proposent d’accéder à ses commandes ou de continuer ses achats.")
para(doc, "Depuis **Mon compte ▸ Mes commandes**, il suit l’avancement de chaque commande via une barre "
          "d’étapes (payée → en préparation → expédiée → livrée) et retrouve le **numéro de suivi** du colis "
          "dès l’expédition.")

h2(doc, "12. L’espace compte")
para(doc, "L’espace **Mon compte** centralise la vie du client :")
bullet(doc, "**Informations personnelles** : nom, téléphone (l’e-mail n’est pas modifiable).")
bullet(doc, "**Historique des commandes** avec statut coloré et accès au détail de chaque commande.")
bullet(doc, "**Points de fidélité** : 1 point est crédité par euro dépensé ; une jauge montre la progression "
            "vers la prochaine récompense.")
bullet(doc, "**Code de parrainage** : le client peut générer un code à partager pour offrir une réduction à "
            "ses proches, et suivre le nombre d’utilisations.")

h2(doc, "13. Demander un retour (côté client)")
para(doc, "Depuis le détail d’une commande **expédiée ou livrée**, le client accède au bouton "
          "**« Retour / Remboursement »**. Le formulaire lui demande :")
numbered(doc, "Un **motif** : rétractation (14 jours), produit défectueux, non conforme à la description, "
              "mauvais article reçu, ou autre.")
numbered(doc, "Une **description** libre (facultative).")
numbered(doc, "Des **photos** (facultatives mais recommandées).")
para(doc, "Avant l’envoi, les **conditions de remboursement** sont rappelées : délai de 14 jours "
          "(rétractation) ou 30 jours (défaut), téléphone réinitialisé en usine et déconnecté du compte "
          "iCloud / Google, IMEI identique à celui expédié, état conforme à l’envoi.")
para(doc, "Une fois la demande envoyée, le client suit son statut (en cours d’examen, approuvée, étiquette "
          "envoyée, colis reçu, inspection, remboursé ou refusé) et retrouve son **numéro RMA** et, le cas "
          "échéant, le **numéro de suivi** de l’étiquette retour et le **montant remboursé**.")
callout(doc, "Le traitement se fait côté administration",
        ["La demande du client **n’est pas automatique** : elle crée un dossier de retour que l’équipe traite "
         "depuis l’espace d’administration (voir §19). C’est l’administrateur qui approuve, envoie l’étiquette, "
         "réceptionne, inspecte et rembourse."],
        kind="info")

h2(doc, "14. Cookies et confidentialité (RGPD)")
para(doc, "À sa première visite, l’internaute voit une **bannière de consentement cookies** conforme aux "
          "recommandations CNIL. Elle propose « Tout accepter », « Refuser » ou « Personnaliser ».")
para(doc, "Le panneau détaillé distingue deux catégories :")
bullet(doc, "**Cookies essentiels** — panier, session, paiement. Toujours actifs, non désactivables.")
bullet(doc, "**Mesure d’audience** — statistiques de visite, désactivable par le visiteur.")
para(doc, "Les outils de mesure d’audience ne se chargent **qu’après** consentement. Le visiteur peut à tout "
          "moment revenir sur son choix via le lien **« Gérer les cookies »** en pied de page.")


# ===========================================================================
# PARTIE III — ESPACE ADMINISTRATION
# ===========================================================================
h1(doc, "Partie III — L’espace d’administration (back-office)")
para(doc, "L’espace d’administration est le poste de pilotage quotidien de l’activité. Cette partie décrit "
          "chaque écran et chaque action, dans l’ordre des tâches courantes : suivre l’activité, traiter les "
          "ventes, gérer les retours et encaissements, entretenir le catalogue et suivre les clients.")

h2(doc, "15. Accès et navigation")
para(doc, "L’administration est accessible via l’adresse **/admin** du site, une fois connecté avec un compte "
          "**administrateur**. Un menu latéral donne accès à toutes les rubriques :")
data_table(
    doc,
    ["Rubrique", "Rôle principal"],
    [
        ["Tableau de bord", "Vue d’ensemble : chiffre d’affaires, commandes à traiter, alertes."],
        ["Commandes", "Traitement des ventes, expéditions, bons fournisseur."],
        ["Retours", "Gestion des demandes de retour et remboursements."],
        ["Produits", "Création et gestion du catalogue et des variantes."],
        ["Prix", "Saisie et mise à jour des prix de vente."],
        ["Marges", "Règles de marge automatiques."],
        ["Paniers", "Suivi des paniers abandonnés et relance."],
        ["Clients", "Fiches clients, segments et historique."],
        ["Litiges", "Suivi des contestations de paiement (chargebacks) Stripe."],
        ["Blocklist", "Liste de blocage anti-fraude."],
        ["Statistiques", "Analyses commerciales et de trafic détaillées."],
    ],
    widths=[1.8, 4.7],
)

h2(doc, "16. Le tableau de bord")
para(doc, "Écran d’accueil de l’administration, il donne une **photographie instantanée** de l’activité :")
bullet(doc, "**Chiffre d’affaires** encaissé, avec comparaison aux 30 jours précédents.")
bullet(doc, "**Commandes à expédier** : nombre de commandes payées en attente de traitement.")
bullet(doc, "**Produits actifs** avec alerte sur les stocks faibles.")
bullet(doc, "**Nombre de clients**.")
bullet(doc, "Un **bandeau paniers abandonnés** (nombre, commandes payées, taux de conversion).")
bullet(doc, "Un **graphique du chiffre d’affaires** sur 30 jours, les **dernières commandes**, les alertes de "
            "**stock faible** et le **top des modèles** vendus.")
para(doc, "Le tableau de bord est une vue de synthèse : les actions se font dans les rubriques dédiées, "
          "accessibles en un clic depuis chaque bloc.")
figure(doc, "06-admin-dashboard.png",
       "Figure 6 — Le tableau de bord : indicateurs clés, paniers, dernières commandes et alertes de stock.")

# ---- 17. Commandes ----
h2(doc, "17. Gérer les ventes et les commandes")
para(doc, "C’est le cœur de l’activité quotidienne. La rubrique **Commandes** liste toutes les ventes et "
          "permet de les faire avancer dans leur cycle de vie.")

h3(doc, "17.1 Le cycle de vie d’une commande")
para(doc, "Chaque commande porte un **statut** qui reflète son avancement. Le tableau ci-dessous récapitule "
          "les statuts et leur signification :")
data_table(
    doc,
    ["Statut", "Signification", "Action attendue de l’équipe"],
    [
        ["**Payée**", "Le client a payé ; la vente est confirmée.", "Regrouper avec les autres ventes pour commander chez le fournisseur."],
        ["**Commande fournisseur**", "La commande a été passée chez le fournisseur.", "Attendre la réception, puis expédier."],
        ["**Expédiée**", "Le colis est parti (IMEI et photos enregistrés).", "Suivre l’acheminement."],
        ["**Livrée**", "Le client a reçu sa commande.", "Aucune — dossier clôturé."],
        ["**Annulée**", "Commande annulée (et remboursée).", "Aucune."],
        ["**Remboursée**", "Remboursement effectué (souvent suite à retour).", "Aucune."],
        ["**Litige**", "Le client conteste le paiement auprès de sa banque.", "Répondre au litige avec les preuves (voir §20)."],
    ],
    widths=[1.7, 2.5, 2.3],
)
callout(doc, "Le fil conducteur",
        ["**Payée → Commande fournisseur → Expédiée → Livrée** est le chemin normal d’une vente réussie. "
         "Les statuts Annulée / Remboursée / Litige sont des branches d’exception."],
        kind="info")

h3(doc, "17.2 Consulter et filtrer les commandes")
para(doc, "La liste des commandes propose des **onglets de filtrage** par statut (Toutes, Payées, Commande "
          "fournisseur, Expédiées, Livrées, Retours, Annulées) et une **recherche** par nom de client, "
          "e-mail, numéro ou identifiant de commande.")
para(doc, "Chaque ligne affiche le numéro de commande, les produits (avec leurs caractéristiques), le client, "
          "le mode de livraison, la date, le statut et le montant. Un clic ouvre le **détail de la commande**.")
figure(doc, "07-admin-commandes.png",
       "Figure 7 — La liste des commandes avec les onglets de statut, la recherche et le bon fournisseur.")

h3(doc, "17.3 Le détail d’une commande")
para(doc, "L’écran de détail présente tout le dossier : articles commandés, récapitulatif financier "
          "(sous-total, remise éventuelle, livraison, total), coordonnées et adresse du client, informations "
          "de livraison et de paiement. Une **frise chronologique** visualise l’avancement.")
para(doc, "Le statut est **modifiable manuellement** via un menu déroulant, mais la plupart des transitions "
          "se font par les **boutons d’action** décrits ci-dessous, qui déclenchent aussi les e-mails et "
          "traitements associés.")
figure(doc, "08-admin-commande-detail.png",
       "Figure 8 — Le détail d’une commande : frise d’avancement, articles, client, livraison et paiement.")

h3(doc, "17.4 Commander chez le fournisseur (bon de commande groupé)")
para(doc, "Plutôt que de commander article par article, l’administration **regroupe toutes les ventes payées** "
          "en un seul bon de commande fournisseur :")
numbered(doc, "Depuis la liste des commandes, cliquer sur **« Commander chez le fournisseur »** "
              "(un badge indique le nombre de commandes concernées).")
numbered(doc, "Le système génère un **bon de commande (Purchase Order)** en brouillon, regroupant les "
              "quantités par variante (même modèle + grade + couleur = une ligne).")
numbered(doc, "Relire le bon, puis l’**imprimer / enregistrer en PDF** pour le transmettre au fournisseur.")
numbered(doc, "Cliquer sur **« Approuver et passer en commande fournisseur »**.")
callout(doc, "Action structurante",
        ["L’approbation du bon fait **basculer d’un coup toutes les commandes « Payées » concernées vers "
         "« Commande fournisseur »**. C’est un point de bascule : vérifiez le bon avant de l’approuver, "
         "car le retour en arrière n’est pas immédiat."],
        kind="danger")

h3(doc, "17.5 Expédier une commande (IMEI + photos)")
para(doc, "À réception des appareils, l’équipe expédie chaque commande. Depuis le détail d’une commande "
          "« Payée » ou « Commande fournisseur », cliquer sur **« Expédier (IMEI + photos) »** :")
numbered(doc, "Saisir l’**IMEI** de chaque appareil (14 à 17 chiffres).")
numbered(doc, "Ajouter au moins une **photo horodatée** de l’appareil / du colis.")
numbered(doc, "Renseigner éventuellement le **numéro et le lien de suivi**.")
numbered(doc, "Valider : la commande passe en **« Expédiée »**, les IMEI et photos sont **verrouillés**, "
              "et un e-mail est envoyé au client.")
callout(doc, "IMEI et photos = protection anti-litige",
        ["Ces preuves (IMEI de l’appareil exact envoyé + photos horodatées) sont **essentielles** en cas de "
         "contestation de paiement ou de retour litigieux. Ne jamais expédier sans les enregistrer."],
        kind="ok")
para(doc, "Si un connecteur transporteur est configuré, un bouton permet aussi de **générer le bordereau "
          "d’expédition** (étiquette PDF) directement depuis la commande, avec récupération automatique du "
          "numéro de suivi.")

h3(doc, "17.6 Marquer comme livrée")
para(doc, "Une fois le colis remis, un clic sur **« Marquer comme livrée »** clôture le dossier "
          "(statut **« Livrée »**).")

h3(doc, "17.7 Annuler et rembourser")
para(doc, "Pour annuler une commande, le bouton **« Annuler + rembourser »** ouvre une fenêtre où l’on saisit "
          "un **message au client** (motif) et le **montant à rembourser** (total par défaut, ajustable pour un "
          "remboursement partiel). La validation déclenche le **remboursement Stripe immédiat**, passe la "
          "commande en **« Annulée »** et informe le client.")

# ---- 18. (merged into 17) keep numbering: next is Retours = 18? we used 19 earlier in callouts. Align. ----
h2(doc, "18. Gérer les retours et remboursements")
para(doc, "La rubrique **Retours** pilote tout le cycle d’un retour produit, depuis la demande du client "
          "jusqu’au remboursement ou au refus. Le workflow est **linéaire** et sécurisé.")
figure(doc, "09-admin-retours.png",
       "Figure 9 — La rubrique Retours et ses onglets de suivi, de « À examiner » à « Remboursé ».")

h3(doc, "18.1 Les statuts d’un retour")
data_table(
    doc,
    ["Statut", "Signification"],
    [
        ["**À examiner** (demandé)", "Le client a soumis une demande ; décision attendue."],
        ["**Approuvé**", "La demande est acceptée ; étiquette retour à envoyer."],
        ["**Étiquette envoyée**", "Le client a reçu son étiquette de retour."],
        ["**Reçu**", "Le colis est arrivé au magasin, pas encore inspecté."],
        ["**En inspection**", "Contrôle de l’appareil en cours."],
        ["**Remboursé**", "Remboursement effectué (dossier clos)."],
        ["**Refusé**", "Demande ou remboursement refusé, avec motif."],
    ],
    widths=[2.1, 4.4],
)

h3(doc, "18.2 Le workflow de traitement, étape par étape")
para(doc, "Depuis le détail d’un retour, l’administration dispose des informations de contrôle : commande "
          "d’origine, **IMEI expédié**, **photos d’expédition**, motif et photos envoyés par le client. "
          "Le traitement suit quatre étapes :")
numbered(doc, "**Décision initiale** — vérifier l’éligibilité (délai, motif, photos) puis **Approuver** ou "
              "**Refuser** (avec motif). Un e-mail est envoyé au client.")
numbered(doc, "**Envoi de l’étiquette** — générer l’étiquette de retour, saisir le numéro de suivi, cliquer "
              "sur **« Étiquette envoyée »**. Le client la reçoit par e-mail.")
numbered(doc, "**Réception du colis** — à l’arrivée, cliquer sur **« Marquer le colis comme reçu »**.")
numbered(doc, "**Inspection** — remplir la check-list de contrôle (ci-dessous) et enregistrer.")
para(doc, "La check-list d’inspection comporte trois contrôles obligatoires :")
bullet(doc, "**IMEI identique** à celui expédié.")
bullet(doc, "**État conforme** aux photos et à l’envoi.")
bullet(doc, "**Réinitialisation usine** effectuée (compte iCloud / Google déconnecté).")
para(doc, "Selon le résultat, le système recommande le remboursement (contrôles conformes) ou le refus "
          "(contrôle non conforme). L’administrateur choisit alors :")
bullet(doc, "**Rembourser via Stripe** — saisir le montant (total ou partiel) et confirmer. Le retour passe "
            "en **« Remboursé »**, le client est notifié.")
bullet(doc, "**Refuser le remboursement** — saisir un motif obligatoire. Le retour passe en **« Refusé »**.")
callout(doc, "Contrôle anti-fraude au retour",
        ["La comparaison de l’**IMEI reçu** avec l’**IMEI expédié** évite la fraude au « swap » (renvoi d’un "
         "appareil différent). En cas de doute, la blocklist (§21) permet de bloquer un client récidiviste."],
        kind="warning")

# ---- 19. Encaissements / paiements / litiges ----
h2(doc, "19. Encaissement et suivi des paiements")
para(doc, "TEL & CASH encaisse les paiements via **Stripe**. Le suivi financier s’organise autour de trois "
          "éléments : la confirmation automatique des paiements, la rubrique **Litiges**, et le tableau de "
          "bord financier.")

h3(doc, "19.1 Comment l’argent est encaissé")
para(doc, "L’encaissement est **automatique** : lorsqu’un client paie sur la page Stripe, Stripe capture le "
          "montant et notifie le site, qui confirme la commande. L’argent est ensuite **reversé sur le compte "
          "bancaire** de l’entreprise selon le calendrier de virement configuré dans Stripe (le tableau de bord "
          "Stripe reste la référence pour les virements et la comptabilité).")
callout(doc, "Où voir les encaissements en détail ?",
        ["Le **tableau de bord Stripe** (dashboard.stripe.com) est la source de vérité pour les paiements, "
         "les virements bancaires, les reçus et la facturation. L’administration TEL & CASH reflète l’état "
         "des commandes ; Stripe gère le flux monétaire."],
        kind="info")

h3(doc, "19.2 Remboursements")
para(doc, "Les remboursements se déclenchent depuis TEL & CASH (annulation de commande §17.7 ou retour §18) "
          "et sont exécutés en direct sur Stripe. Le montant apparaît sur la carte du client sous 5 à 10 jours "
          "ouvrés. Un remboursement peut être **total ou partiel**.")

h3(doc, "19.3 Litiges et contestations (chargebacks)")
para(doc, "Quand un client conteste un paiement auprès de sa banque, Stripe ouvre un **litige** (chargeback). "
          "La rubrique **Litiges** de l’administration liste ces contestations avec leur motif, leur montant, "
          "le client, la commande liée et le statut Stripe.")
para(doc, "Depuis un litige, l’administrateur ouvre la **commande concernée** pour rassembler les preuves "
          "(IMEI, photos d’expédition, suivi de livraison) à soumettre. La **réponse au litige elle-même se "
          "fait dans le tableau de bord Stripe**, qui pilote la procédure et les délais.")
callout(doc, "Réagir vite",
        ["Les litiges ont un **délai de réponse court** imposé par les réseaux bancaires. Traitez-les en "
         "priorité et joignez systématiquement les preuves d’expédition enregistrées à la commande."],
        kind="danger")
figure(doc, "15-admin-litiges.png",
       "Figure 15 — La rubrique Litiges : suivi des contestations de paiement (chargebacks) Stripe.")

# ---- 20. Catalogue produits ----
h2(doc, "20. Gérer le catalogue produits")
para(doc, "La rubrique **Produits** permet de créer, organiser et entretenir le catalogue. Il est structuré "
          "en quatre onglets selon la source et la catégorie :")
bullet(doc, "**Téléphones boutique** et **Accessoires boutique** — saisis manuellement par l’équipe.")
bullet(doc, "**Téléphones Fluxitron** et **Accessoires Fluxitron** — issus de la synchronisation fournisseur "
            "(voir annexe technique) ; l’équipe en contrôle seulement la **visibilité** (activer/désactiver).")

h3(doc, "20.1 Vue d’ensemble et actions groupées")
para(doc, "Les téléphones s’affichent en **arborescence** (modèle ▸ stockage ▸ grade ▸ couleurs) ou à plat. "
          "On filtre par marque, on trie par modèle, stock ou prix, et on recherche par nom, marque, modèle "
          "ou référence.")
para(doc, "Des **actions groupées** (après sélection par cases à cocher) permettent d’**activer**, "
          "**désactiver** ou **supprimer** plusieurs produits d’un coup. La désactivation retire un produit "
          "du catalogue client sans le supprimer.")
figure(doc, "10-admin-produits.png",
       "Figure 10 — Le catalogue en arborescence (modèle ▸ stockage ▸ grade ▸ couleurs) et les actions groupées.")

h3(doc, "20.2 Créer un nouveau produit")
para(doc, "Le formulaire de création (**Produits ▸ Nouveau**) se remplit en plusieurs sections :")
numbered(doc, "**Identité** — marque, modèle, catégorie.")
numbered(doc, "**Spécifications** — caractéristiques techniques (pré-remplies pour les iPhone) et description.")
numbered(doc, "**Déclinaisons** — cocher les stockages, grades (A, B, C) et couleurs proposés.")
numbered(doc, "**Grille de prix** — saisir un prix par combinaison stockage × grade.")
numbered(doc, "**Photos par couleur** — au moins une photo par couleur choisie.")
para(doc, "À l’enregistrement, le système crée automatiquement **une variante par combinaison** "
          "(stockage × grade × couleur).")
callout(doc, "Un prix vide = variante non vendable",
        ["Une variante sans prix (ou à 0) est **grisée** et non commandable côté client. Pour rendre une "
         "variante disponible, il suffit de lui attribuer un prix."],
        kind="info")
figure(doc, "16-admin-produit-nouveau.png",
       "Figure 11 — Le formulaire de création d’un produit : identité, déclinaisons, grille de prix et photos.")

h3(doc, "20.3 Modifier un produit")
para(doc, "L’édition d’un produit permet d’ajuster les specs, les prix, les photos et l’activation par "
          "variante. La **suppression** est possible si le produit n’est lié à aucune commande passée ; "
          "sinon, on privilégie la **désactivation** afin de préserver l’historique des ventes.")

h2(doc, "21. Gérer les prix et les marges")
h3(doc, "21.1 Rubrique Prix")
para(doc, "La rubrique **Prix** centralise la saisie des prix de vente par modèle. Chaque modèle se déplie en "
          "une grille (lignes = stockages, colonnes = grades A/B/C). On saisit les prix, on clique sur "
          "**Appliquer**, et l’on peut activer un **prix barré** (promotion) par ligne. Un interrupteur "
          "**active ou désactive** toutes les variantes d’un modèle en une fois.")
figure(doc, "11-admin-prix.png",
       "Figure 12 — La rubrique Prix : grille stockage × grade par modèle, avec activation et promotions.")
h3(doc, "21.2 Rubrique Marges")
para(doc, "La rubrique **Marges** permet de définir des **règles de marge automatiques** (par marque, modèle "
          "ou produit), en pourcentage ou en montant fixe, avec des règles d’arrondi (au centime, à l’euro, "
          "prix finissant par .99…). Elle sert à **standardiser la rentabilité** et éviter les oublis. "
          "Un aperçu montre l’effet des règles avant application.")

h2(doc, "22. Paniers abandonnés et relance")
para(doc, "La rubrique **Paniers** suit les commandes **initiées mais non payées** (le client a lancé le "
          "paiement sans le finaliser). Elle affiche le nombre de paniers, le nombre de commandes payées et "
          "le **taux de conversion**, ainsi que le détail de chaque panier (articles, client, montant "
          "potentiel, date).")
callout(doc, "Relance automatique par e-mail",
        ["Une **relance automatique** est envoyée par e-mail aux clients ayant abandonné leur panier "
         "(une seule relance par panier, via une tâche planifiée quotidienne). Cette rubrique permet d’en "
         "suivre l’efficacité et d’identifier d’éventuels points de friction au paiement."],
        kind="info")
figure(doc, "13-admin-paniers.png",
       "Figure 13 — Le suivi des paniers abandonnés et du taux de conversion.")

h2(doc, "23. Suivre les clients (dossiers)")
para(doc, "La rubrique **Clients** regroupe la base clients. La liste offre une **recherche** (nom, e-mail, "
          "téléphone) et des **segments** automatiques calculés sur le total dépensé :")
bullet(doc, "**VIP** — plus de 500 € dépensés.")
bullet(doc, "**Fidèle** — entre 250 € et 500 €.")
bullet(doc, "**Nouveau** — moins de 250 € ou une seule commande.")
para(doc, "La **fiche client** affiche les indicateurs clés (total dépensé, nombre de commandes, panier "
          "moyen), les coordonnées, la dernière adresse de livraison, une zone de **notes internes** et "
          "l’**historique complet des commandes** (cliquables vers leur détail). C’est l’outil central pour "
          "le suivi d’un dossier client au téléphone ou par e-mail.")
figure(doc, "12-admin-clients.png",
       "Figure 14 — La base clients : recherche, segments (VIP, Fidèle, Nouveau) et indicateurs par client.")

h2(doc, "24. Blocklist anti-fraude")
para(doc, "La rubrique **Blocklist** permet de **bloquer** des acheteurs à risque. On y ajoute une entrée en "
          "précisant son **type** (e-mail, IP, IMEI, téléphone, identifiant client), sa **valeur** et un "
          "**motif** (ex. « chargeback frauduleux », « swap au retour »). Les entrées peuvent être supprimées. "
          "Le motif constitue une **trace** utile en cas de litige.")

h2(doc, "25. Statistiques avancées")
para(doc, "La rubrique **Statistiques** propose des analyses sur 7, 30, 90 jours ou 12 mois :")
bullet(doc, "**Commerce** — chiffre d’affaires, nombre de commandes, panier moyen, nouveaux clients, avec "
            "comparaison à la période précédente.")
bullet(doc, "**Top produits** par chiffre d’affaires.")
bullet(doc, "**Trafic** (si la mesure d’audience est active) — visiteurs, pages vues, taux de conversion, "
            "sources de trafic, répartition par appareil, pages les plus vues.")
para(doc, "Ces données sont **consultatives** : elles éclairent les décisions commerciales sans déclencher "
          "d’action.")
figure(doc, "14-admin-statistiques.png",
       "Figure 16 — Les statistiques avancées : commerce, top produits et trafic sur la période choisie.")


# ===========================================================================
# PARTIE IV — OPÉRATIONS QUOTIDIENNES
# ===========================================================================
h1(doc, "Partie IV — Opérations quotidiennes & check-lists")
para(doc, "Cette partie propose des routines prêtes à l’emploi pour l’équipe. Elle synthétise les procédures "
          "des parties précédentes sous forme de mémos actionnables.")

h2(doc, "26. Routine quotidienne recommandée")
numbered(doc, "Ouvrir le **tableau de bord** : repérer les commandes à expédier et les alertes.")
numbered(doc, "Traiter les **nouvelles ventes payées** : générer / approuver le **bon de commande fournisseur**.")
numbered(doc, "**Expédier** les commandes dont les appareils sont réceptionnés (IMEI + photos).")
numbered(doc, "Passer en revue la rubrique **Retours** : approuver, envoyer les étiquettes, inspecter, rembourser.")
numbered(doc, "Vérifier la rubrique **Litiges** : répondre en priorité à tout nouveau litige.")
numbered(doc, "Contrôler les **paniers abandonnés** et l’état général de l’activité.")

h2(doc, "27. Check-list : expédition sécurisée")
callout(doc, "À vérifier avant chaque expédition",
        ["☐ IMEI de **l’appareil réellement envoyé** saisi et vérifié.",
         "☐ Au moins une **photo horodatée** de l’appareil et du colis.",
         "☐ Appareil **testé** et **réinitialisé** si nécessaire.",
         "☐ **Numéro de suivi** renseigné.",
         "☐ Bordereau collé, colis correctement protégé."],
        kind="ok")

h2(doc, "28. Check-list : traitement d’un retour")
numbered(doc, "Vérifier l’**éligibilité** (délai, motif) → Approuver ou Refuser.")
numbered(doc, "Envoyer l’**étiquette** et enregistrer le suivi.")
numbered(doc, "À réception, **marquer le colis reçu**.")
numbered(doc, "**Inspecter** : IMEI identique, état conforme, réinitialisation usine.")
numbered(doc, "**Rembourser** (total ou partiel) ou **refuser** avec motif.")

h2(doc, "29. En cas de litige de paiement")
numbered(doc, "Ouvrir la rubrique **Litiges** et identifier la commande concernée.")
numbered(doc, "Rassembler les **preuves** : IMEI, photos d’expédition, preuve de livraison / suivi.")
numbered(doc, "Soumettre la réponse et les preuves **dans le tableau de bord Stripe**, avant l’échéance.")
numbered(doc, "Si le client est manifestement de mauvaise foi, l’ajouter à la **blocklist**.")

h2(doc, "30. Points de vigilance")
callout(doc, "Actions sensibles",
        ["**Approbation d’un bon fournisseur** : bascule en masse de commandes — vérifier avant de valider.",
         "**Suppression d’un produit** lié à des commandes : préférer la **désactivation** pour garder l’historique.",
         "**Remboursement** : opération irréversible côté Stripe — contrôler le montant.",
         "**Ne jamais expédier** sans IMEI ni photo : c’est la seule protection en cas de litige."],
        kind="danger")


# ===========================================================================
# PARTIE V — ANNEXE TECHNIQUE
# ===========================================================================
h1(doc, "Partie V — Annexe technique")
para(doc, "Cette annexe s’adresse au responsable technique. Elle documente l’architecture réelle de la "
          "plateforme pour en assurer la maintenance et, si besoin, la reprise.")

h2(doc, "31. Pile technique (stack)")
data_table(
    doc,
    ["Technologie", "Rôle"],
    [
        ["**Next.js 15** (App Router)", "Framework full-stack : pages, composants serveur/client, routes API."],
        ["**React 19**", "Interface utilisateur."],
        ["**TypeScript**", "Typage statique (front et back)."],
        ["**Supabase**", "Base de données PostgreSQL + Authentification + Stockage de fichiers + sécurité RLS."],
        ["**Stripe**", "Paiements et remboursements, webhooks, facturation."],
        ["**Zustand**", "État du panier côté client (persistant)."],
        ["**Tailwind CSS** + **Framer Motion**", "Style et animations."],
        ["**Nodemailer / Resend**", "E-mails transactionnels (confirmations, relances)."],
    ],
    widths=[2.4, 4.1],
)
para(doc, "**Scripts npm** principaux : `npm run dev` (développement), `npm run build` (build de production), "
          "`npm run start` (serveur de production), `npm run lint` (analyse du code). Des scripts dédiés "
          "gèrent le catalogue et la synchronisation fournisseur (voir §36).")

h2(doc, "32. Hébergement et déploiement")
para(doc, "Le site est hébergé sur **Vercel** (plateforme optimisée pour Next.js). Le déploiement est "
          "automatique à chaque mise à jour du code. Vercel héberge également les **tâches planifiées** "
          "(crons) décrites au §37. Les images produits sont servies depuis le **stockage Supabase**.")

h2(doc, "33. Base de données et sécurité")
para(doc, "La base **PostgreSQL** est gérée par Supabase, avec la **sécurité au niveau des lignes** "
          "(Row-Level Security) : chaque utilisateur n’accède qu’à ses propres données ; les administrateurs "
          "disposent d’un accès étendu. Les principales tables :")
data_table(
    doc,
    ["Table", "Contenu"],
    [
        ["`profiles`", "Profils utilisateurs (nom, e-mail, téléphone, **rôle** admin/client)."],
        ["`products`", "Catalogue : marque, modèle, stockage, couleur, grade, prix, coût, images, stock, source, référence."],
        ["`cart_items`", "Paniers des clients connectés."],
        ["`orders`", "Commandes : montant, statut, adresse, suivi, numéro de commande, session Stripe."],
        ["`order_items`", "Lignes de commande, avec prix et coût **figés** au moment de l’achat."],
        ["`loyalty_points`", "Points de fidélité (1 point par euro)."],
        ["`referral_codes`", "Codes de parrainage et leurs utilisations."],
        ["`return_requests`", "Demandes de retour (RMA), statut, inspection, remboursement."],
        ["`disputes`", "Litiges Stripe (chargebacks)."],
        ["`margin_rules` / `margin_settings`", "Règles et paramètres de marge."],
        ["`supplier_sync_settings`", "Réglages de synchronisation du stock fournisseur."],
        ["`stripe_events`", "Journal d’idempotence des webhooks Stripe."],
    ],
    widths=[2.3, 4.2],
)

h2(doc, "34. Authentification et rôles")
para(doc, "L’authentification repose sur **Supabase Auth**. À l’inscription, un profil est créé "
          "automatiquement dans la table `profiles`. Le **rôle** (client ou administrateur) y est stocké. "
          "Un **middleware** protège toutes les pages et API `/admin/*` : il rafraîchit la session à chaque "
          "requête et refuse l’accès aux non-administrateurs.")
para(doc, "Trois clients d’accès à la base coexistent, selon le contexte :")
bullet(doc, "**Client navigateur** — opérations publiques, soumis à la sécurité RLS.")
bullet(doc, "**Client serveur** — composants serveur et API, avec la session de l’utilisateur.")
bullet(doc, "**Client admin** (clé de service) — utilisé uniquement côté serveur (webhooks, crons, API admin), "
            "contourne la RLS. Cette clé ne doit **jamais** être exposée côté client.")
callout(doc, "Attribuer le rôle administrateur",
        ["Pour donner l’accès admin à un compte, on positionne son champ `role` à `admin` dans la table "
         "`profiles` (via le tableau de bord Supabase). À manipuler avec précaution."],
        kind="warning")

h2(doc, "35. Paiements Stripe (flux technique)")
para(doc, "Le parcours de paiement s’appuie sur **Stripe Checkout** (page hébergée par Stripe) :")
numbered(doc, "Le site crée une **session de paiement** (`/api/checkout`) après validation du panier et "
              "contrôles anti-fraude, et enregistre une pré-commande au statut « en attente ».")
numbered(doc, "Le client paie sur la page Stripe.")
numbered(doc, "Stripe appelle le **webhook** du site (`/api/webhooks/stripe`), authentifié par signature.")
numbered(doc, "Le webhook confirme la commande, **décrémente le stock de façon atomique** (avec remboursement "
              "automatique en cas de survente), attribue les points de fidélité, vide le panier et déclenche "
              "les e-mails.")
para(doc, "Les événements traités incluent la réussite/échec du paiement, l’expiration de session, les "
          "remboursements et l’ouverture / mise à jour des litiges. Un journal d’idempotence "
          "(`stripe_events`) évite tout double traitement.")

h2(doc, "36. Intégration Fluxitron (stock fournisseur)")
para(doc, "**Fluxitron** (Hub fournisseur) sert de **miroir de stock**. Un import quotidien récupère le "
          "**Catalog Feed** du fournisseur, le rapproche du catalogue boutique et met à jour la disponibilité. "
          "Ce mécanisme permet, en option, de **griser** automatiquement les variantes sans stock fournisseur.")
para(doc, "Une API dédiée (`/api/v1/…`, authentifiée par clé `X-Api-Key`) expose produits, prix, stock, "
          "commandes et catégories au Hub. Des **webhooks** informent le Hub de la création, mise à jour ou "
          "annulation des commandes. La spécification `public/openapi.yaml` sert à configurer le connecteur.")
callout(doc, "Le prix reste maîtrisé par la boutique",
        ["Fluxitron alimente **uniquement le signal de disponibilité** (stock). Les **prix de vente restent "
         "gérés manuellement** par l’équipe via les rubriques Prix et Marges."],
        kind="info")

h2(doc, "37. Tâches planifiées (crons)")
data_table(
    doc,
    ["Tâche", "Fréquence", "Rôle"],
    [
        ["Synchronisation fournisseur", "Quotidienne (matin)", "Importe le Catalog Feed Fluxitron et met à jour la disponibilité."],
        ["Relance paniers abandonnés", "Quotidienne", "Envoie un e-mail de relance aux paniers non finalisés (une seule fois par panier)."],
    ],
    widths=[2.2, 1.6, 2.7],
)
para(doc, "Ces tâches sont déclenchées par Vercel et protégées par un secret (`CRON_SECRET`). Elles peuvent "
          "aussi être lancées manuellement par un administrateur.")

h2(doc, "38. Variables d’environnement")
para(doc, "La configuration sensible est fournie par variables d’environnement (jamais dans le code). "
          "Référence complète :")
h3(doc, "Supabase")
data_table(doc, ["Variable", "Description"], [
    ["`NEXT_PUBLIC_SUPABASE_URL`", "URL du projet Supabase."],
    ["`NEXT_PUBLIC_SUPABASE_ANON_KEY`", "Clé publique (navigateur), soumise à la RLS."],
    ["`SUPABASE_SERVICE_ROLE_KEY`", "Clé de service (serveur uniquement) — contourne la RLS. **Secret.**"],
], widths=[2.9, 3.6])
h3(doc, "Stripe")
data_table(doc, ["Variable", "Description"], [
    ["`STRIPE_SECRET_KEY`", "Clé API secrète Stripe."],
    ["`NEXT_PUBLIC_STRIPE_PUBLISHABLE_KEY`", "Clé publique Stripe (front)."],
    ["`STRIPE_WEBHOOK_SECRET`", "Secret de validation de signature du webhook."],
], widths=[2.9, 3.6])
h3(doc, "Fluxitron")
data_table(doc, ["Variable", "Description"], [
    ["`FLUXITRON_API_KEY`", "Clé du connecteur (API `/api/v1`)."],
    ["`FLUXITRON_WEBHOOK_URL`", "URL du webhook du Hub."],
    ["`FLUXITRON_WEBHOOK_API_KEY`", "Clé d’authentification des webhooks sortants."],
    ["`FLUXITRON_FEED_URL`", "URL du Catalog Feed (JSON)."],
    ["`FLUXITRON_FEED_TOKEN`", "Jeton d’accès au Catalog Feed."],
], widths=[2.9, 3.6])
h3(doc, "Application, e-mails & divers")
data_table(doc, ["Variable", "Description"], [
    ["`NEXT_PUBLIC_APP_URL`", "URL publique du site."],
    ["`NEXT_PUBLIC_SHIPPING_FEE_EUR`", "Frais de livraison (défaut 9,90 €)."],
    ["`CRON_SECRET`", "Secret d’authentification des tâches planifiées Vercel."],
    ["`RESEND_API_KEY` / `SMTP_*`", "Envoi des e-mails transactionnels (Resend ou SMTP)."],
    ["`MERCHANT_EMAIL`", "Adresse recevant les alertes « nouvelle commande »."],
    ["`BOXTAL_*`", "Connecteur transporteur (bordereaux d’expédition), le cas échéant."],
    ["`UMAMI_*`", "Mesure d’audience externe (optionnelle)."],
], widths=[2.9, 3.6])

h2(doc, "39. Organisation du code")
para(doc, "Le code applicatif est regroupé sous `src/` :")
data_table(doc, ["Dossier", "Contenu"], [
    ["`app/`", "Pages et routes API (Next.js App Router), y compris `admin/`, `api/` et les pages publiques."],
    ["`components/`", "Composants d’interface réutilisables."],
    ["`contexts/`", "Contextes React (authentification, notifications)."],
    ["`store/`", "États globaux côté client (panier, consentement cookies)."],
    ["`lib/`", "Logique métier et connecteurs : Supabase, Stripe, e-mails, livraison, fournisseur, anti-fraude."],
    ["`data/`", "Données statiques (marques, modèles)."],
    ["`middleware.ts`", "Protection des routes et rafraîchissement de session."],
], widths=[1.7, 4.8])

h2(doc, "40. Maintenance : points d’attention")
bullet(doc, "**Achat invité désactivé** : le paiement requiert un compte (une évolution est préparée mais "
            "non activée).")
bullet(doc, "**Grisage fournisseur** : conçu « fail-open » — en cas de données manquantes ou périmées, le "
            "produit reste vendable.")
bullet(doc, "**E-mails non bloquants** : une défaillance d’envoi n’interrompt jamais le paiement ni les crons.")
bullet(doc, "**Webhooks Fluxitron** : envoyés en tâche de fond avec réessais ; sans impact sur le flux "
            "principal en cas d’échec.")
bullet(doc, "**Tests** : quelques tests unitaires (Vitest) ; pas de suite end-to-end. À renforcer en amont "
            "d’évolutions majeures.")


# ===========================================================================
# GLOSSAIRE
# ===========================================================================
h1(doc, "Glossaire")
data_table(
    doc,
    ["Terme", "Définition"],
    [
        ["**Grade (A/B/C)**", "Classement de l’état esthétique d’un téléphone reconditionné."],
        ["**Variante**", "Combinaison unique stockage × grade × couleur d’un modèle, avec son prix et ses photos."],
        ["**Sell-to-order**", "Vente à la commande : approvisionnement fournisseur déclenché par la vente ; pas de stock affiché."],
        ["**IMEI**", "Numéro d’identification unique d’un téléphone (15 chiffres) — preuve de l’appareil exact expédié."],
        ["**RMA**", "Numéro unique attribué à une demande de retour (Return Merchandise Authorization)."],
        ["**Bon de commande fournisseur**", "Document regroupant les ventes payées pour commander les appareils en une fois."],
        ["**Chargeback / Litige**", "Contestation d’un paiement par le client auprès de sa banque, gérée via Stripe."],
        ["**Webhook**", "Notification automatique envoyée de service à service (ex. Stripe → site après paiement)."],
        ["**Fluxitron**", "Hub fournisseur alimentant le signal de stock du catalogue."],
        ["**RLS**", "Row-Level Security : règles de sécurité au niveau des lignes de la base de données."],
        ["**Cron**", "Tâche planifiée exécutée automatiquement à intervalle régulier."],
    ],
    widths=[2.2, 4.3],
)

spacer(doc, 14)
end = doc.add_paragraph()
add_bottom_rule(end, LINE, size=8)
re = end.add_run("Fin du manuel — TEL & CASH · Manuel d’utilisation v1.0 · 7 juillet 2026")
re.italic = True
re.font.size = Pt(9)
re.font.color.rgb = C(GRAY)


# ---- Génération du sommaire à l'ouverture ----
enable_update_fields(doc)

OUT = "/Users/fantin/Documents/CODE/Site WEB/TEL and CASH/docs/Manuel_Utilisation_TEL_and_CASH.docx"
doc.save(OUT)
print("OK ->", OUT)
